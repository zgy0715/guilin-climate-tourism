#!/usr/bin/env python3
"""
Markdown to Word converter - 严格按照论文格式要求
- 论文题目：三号黑体(16pt)，居中
- 摘要标题：四号黑体(14pt)，居中
- 正文：小四宋体(12pt)
- 英文：小四Times New Roman(12pt)
- 行距：单倍行距（固定值20磅），段前段后0.5行
- 数学公式：单倍行距，居中，公式编辑器格式
"""

import re
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


def set_run_font(run, cn_font='宋体', en_font='Times New Roman', size=12, bold=False):
    """设置run的中英文字体、字号、加粗"""
    run.font.name = en_font
    run._element.rPr.rFonts.set(qn('w:eastAsia'), cn_font)
    run.font.size = Pt(size)
    run.bold = bold


def set_paragraph_format(p, line_spacing=20, space_before=10, space_after=10,
                          alignment=None, first_line_indent=None):
    """设置段落格式：行距、段前段后、对齐、首行缩进"""
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(line_spacing)
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if alignment is not None:
        p.alignment = alignment
    if first_line_indent is not None:
        pf.first_line_indent = Cm(first_line_indent)


def add_heading_custom(doc, text, cn_font='黑体', en_font='Times New Roman', size=16,
                        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=12, space_after=6):
    """添加标题（三号黑体居中）"""
    p = doc.add_paragraph()
    set_paragraph_format(p, line_spacing=20, space_before=space_before,
                          space_after=space_after, alignment=alignment)
    run = p.add_run(text)
    set_run_font(run, cn_font=cn_font, en_font=en_font, size=size, bold=True)
    return p


def add_body_text(doc, text, cn_font='宋体', en_font='Times New Roman', size=12,
                   alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_indent=0.74):
    """添加正文段落（小四宋体，两端对齐，首行缩进2字符）"""
    p = doc.add_paragraph()
    set_paragraph_format(p, line_spacing=20, space_before=10, space_after=10,
                          alignment=alignment, first_line_indent=first_line_indent)

    # 处理加粗标记 **text**
    parts = re.split(r'\*\*(.*?)\*\*', text)
    for i, part in enumerate(parts):
        if part:
            run = p.add_run(part)
            is_bold = (i % 2 == 1)
            set_run_font(run, cn_font=cn_font, en_font=en_font, size=size, bold=is_bold)
    return p


def add_math_paragraph(doc, text, tag_num=None):
    """添加数学公式（居中，公式编辑器格式，段前段后0.5行）"""
    p = doc.add_paragraph()
    set_paragraph_format(p, line_spacing=20, space_before=10, space_after=10,
                          alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # 清理LaTeX并转换为可读文本
    display_text = text
    display_text = display_text.replace('\\frac', ' ')
    display_text = display_text.replace('\\exp', 'exp')
    display_text = display_text.replace('\\text{', '')
    display_text = display_text.replace('\\', '')
    display_text = display_text.replace('{', '(')
    display_text = display_text.replace('}', ')')
    display_text = display_text.replace('^', 'ⁿ')
    display_text = display_text.replace('_', 'ₙ')
    display_text = display_text.replace('\\sum', '∑')
    display_text = display_text.replace('\\sqrt', '√')
    display_text = display_text.replace('\\ln', 'ln')
    display_text = display_text.replace('\\max', 'max')
    display_text = display_text.replace('\\min', 'min')
    display_text = display_text.replace('\\mid', '|')
    display_text = display_text.replace('\\forall', '∀')
    display_text = display_text.replace('\\median', 'median')
    display_text = display_text.replace('\\cdot', '·')
    display_text = display_text.replace('\\times', '×')

    run = p.add_run(display_text)
    set_run_font(run, cn_font='宋体', en_font='Times New Roman', size=11)

    if tag_num:
        run2 = p.add_run(f'    ({tag_num})')
        set_run_font(run2, cn_font='宋体', en_font='Times New Roman', size=11)

    return p


def set_cell_font(cell, cn_font='宋体', en_font='Times New Roman', size=10.5):
    """设置表格单元格字体"""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            set_run_font(run, cn_font=cn_font, en_font=en_font, size=size)


def parse_table(lines):
    """解析markdown表格"""
    rows = []
    for line in lines:
        line = line.strip()
        if line.startswith('|') and line.endswith('|'):
            cells = [c.strip() for c in line.split('|')[1:-1]]
            if all(set(c) <= set('-: ') for c in cells):
                continue
            rows.append(cells)
    return rows


def add_table(doc, rows):
    """添加表格（三线表格式）"""
    if not rows:
        return

    num_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            if j < num_cols:
                cell = table.cell(i, j)
                cell.text = cell_text
                set_cell_font(cell, '宋体', 'Times New Roman', 10.5)
                # 表头加粗
                if i == 0:
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.bold = True
    return table


def add_image(doc, img_path, caption=None):
    """添加图片（居中，图片下方图注）"""
    if os.path.exists(img_path):
        p = doc.add_paragraph()
        set_paragraph_format(p, line_spacing=20, space_before=10, space_after=6,
                              alignment=WD_ALIGN_PARAGRAPH.CENTER)
        run = p.add_run()
        try:
            run.add_picture(img_path, width=Inches(5.5))
        except Exception:
            run.add_text(f'[图片: {os.path.basename(img_path)}]')

        if caption:
            cap_p = doc.add_paragraph()
            set_paragraph_format(cap_p, line_spacing=20, space_before=2, space_after=10,
                                  alignment=WD_ALIGN_PARAGRAPH.CENTER)
            cap_run = cap_p.add_run(caption)
            set_run_font(cap_run, cn_font='宋体', en_font='Times New Roman', size=10.5, bold=True)
    else:
        p = doc.add_paragraph()
        set_paragraph_format(p, line_spacing=20, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        run = p.add_run(f'[图片: {os.path.basename(img_path)} - 需手动插入]')
        set_run_font(run, cn_font='宋体', en_font='Times New Roman', size=10.5)
        run.font.color.rgb = RGBColor(128, 128, 128)


def add_list_item(doc, text, style='List Bullet'):
    """添加列表项"""
    p = doc.add_paragraph(style=style)
    set_paragraph_format(p, line_spacing=20, space_before=10, space_after=10)
    parts = re.split(r'\*\*(.*?)\*\*', text)
    for i, part in enumerate(parts):
        if part:
            run = p.add_run(part)
            is_bold = (i % 2 == 1)
            set_run_font(run, cn_font='宋体', en_font='Times New Roman', size=12, bold=is_bold)
    return p


def add_numbered_item(doc, num, text):
    """添加编号列表项（小四宋体，首行缩进）"""
    p = doc.add_paragraph()
    set_paragraph_format(p, line_spacing=20, space_before=10, space_after=10,
                          first_line_indent=0.74)
    full_text = f'（{num}）{text}'
    parts = re.split(r'\*\*(.*?)\*\*', full_text)
    for i, part in enumerate(parts):
        if part:
            run = p.add_run(part)
            is_bold = (i % 2 == 1)
            set_run_font(run, cn_font='宋体', en_font='Times New Roman', size=12, bold=is_bold)
    return p


def add_ref_item(doc, text):
    """添加参考文献条目（悬挂缩进）"""
    p = doc.add_paragraph()
    set_paragraph_format(p, line_spacing=20, space_before=6, space_after=6)
    p.paragraph_format.first_line_indent = Cm(-0.74)
    p.paragraph_format.left_indent = Cm(0.74)
    run = p.add_run(text)
    set_run_font(run, cn_font='宋体', en_font='Times New Roman', size=10.5)
    return p


def md_to_docx(md_path, docx_path):
    """主转换函数"""
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    base_dir = os.path.dirname(os.path.dirname(os.path.abspath(md_path)))  # 项目根目录

    doc = Document()

    # ===== 设置页面边距（A4，上下2.54cm，左右3.18cm）=====
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.18)
        section.right_margin = Cm(3.18)

    # ===== 设置默认样式 =====
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    style.paragraph_format.line_spacing = Pt(20)

    i = 0
    table_buffer = []

    while i < len(lines):
        line = lines[i].rstrip('\n')

        # 空行：刷新表格
        if not line.strip():
            if table_buffer:
                rows = parse_table(table_buffer)
                add_table(doc, rows)
                table_buffer = []
            i += 1
            continue

        # 表格行
        if line.strip().startswith('|') and line.strip().endswith('|'):
            table_buffer.append(line)
            i += 1
            continue

        # ===== 图片引用 =====
        img_match = re.match(r'^!\[(.*?)\]\((.*?)\)', line.strip())
        if img_match:
            img_path = img_match.group(2)
            if not os.path.isabs(img_path):
                img_path = os.path.join(base_dir, img_path)
            caption = None
            if i + 1 < len(lines):
                next_line = lines[i + 1].strip()
                cap_match = re.match(r'<center>(.*?)</center>', next_line)
                if cap_match:
                    caption = cap_match.group(1)
                    i += 1
            add_image(doc, img_path, caption)
            i += 1
            continue

        # 跳过center标签
        if line.strip().startswith('<center>') and line.strip().endswith('</center>'):
            i += 1
            continue

        # ===== 数学公式 =====
        if line.strip().startswith('$$'):
            math_lines = [line.strip().replace('$$', '')]
            if not line.strip().endswith('$$') or line.strip() == '$$':
                i += 1
                while i < len(lines):
                    ml = lines[i].rstrip('\n').strip()
                    if ml.endswith('$$'):
                        math_lines.append(ml.replace('$$', ''))
                        break
                    math_lines.append(ml)
                    i += 1

            math_text = ' '.join(math_lines).strip()
            tag_match = re.search(r'\\tag\{(\d+)\}', math_text)
            tag_num = tag_match.group(1) if tag_match else None
            math_text = re.sub(r'\\tag\{\d+\}', '', math_text).strip()
            add_math_paragraph(doc, math_text, tag_num)
            i += 1
            continue

        # ===== 一级标题（# 论文题目）：三号黑体居中 =====
        if line.startswith('# ') and not line.startswith('## '):
            text = line[2:].strip()
            add_heading_custom(doc, text, cn_font='黑体', en_font='Times New Roman',
                                size=16, space_before=12, space_after=6)
            i += 1
            continue

        # ===== 二级标题（## 摘要、## 一、问题重述）：四号黑体居中 =====
        if line.startswith('## '):
            text = line[3:].strip()
            add_heading_custom(doc, text, cn_font='黑体', en_font='Times New Roman',
                                size=14, space_before=12, space_after=6)
            i += 1
            continue

        # ===== 三级标题（### 5.1）：小四黑体 =====
        if line.startswith('### '):
            text = line[4:].strip()
            p = doc.add_paragraph()
            set_paragraph_format(p, line_spacing=20, space_before=12, space_after=6)
            run = p.add_run(text)
            set_run_font(run, cn_font='黑体', en_font='Times New Roman', size=12, bold=True)
            i += 1
            continue

        # ===== 四级标题（#### 5.1.1）：小四黑体 =====
        if line.startswith('#### '):
            text = line[5:].strip()
            p = doc.add_paragraph()
            set_paragraph_format(p, line_spacing=20, space_before=10, space_after=6)
            run = p.add_run(text)
            set_run_font(run, cn_font='黑体', en_font='Times New Roman', size=12, bold=True)
            i += 1
            continue

        # ===== 分隔线 =====
        if line.strip() == '---':
            p = doc.add_paragraph()
            set_paragraph_format(p, line_spacing=20, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            run = p.add_run('─' * 40)
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(192, 192, 192)
            i += 1
            continue

        # ===== 列表项 =====
        if re.match(r'^[\-\*] ', line):
            text = line[2:].strip()
            add_list_item(doc, text, 'List Bullet')
            i += 1
            continue

        # ===== 编号列表（（1）） =====
        num_match = re.match(r'^（(\d+)）', line)
        if num_match:
            text = line[num_match.end():].strip()
            add_numbered_item(doc, num_match.group(1), text)
            i += 1
            continue

        # ===== 数字编号列表 =====
        num_match2 = re.match(r'^(\d+)\.\s', line)
        if num_match2:
            text = line[num_match2.end():].strip()
            p = doc.add_paragraph(style='List Number')
            set_paragraph_format(p, line_spacing=20, space_before=10, space_after=10)
            parts = re.split(r'\*\*(.*?)\*\*', text)
            for j, part in enumerate(parts):
                if part:
                    run = p.add_run(part)
                    is_bold = (j % 2 == 1)
                    set_run_font(run, cn_font='宋体', en_font='Times New Roman', size=12, bold=is_bold)
            i += 1
            continue

        # ===== 关键词 =====
        if line.startswith('**关键词：**'):
            text = line.replace('**关键词：**', '').strip()
            p = doc.add_paragraph()
            set_paragraph_format(p, line_spacing=20, space_before=10, space_after=10,
                                  alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_indent=0.74)
            run1 = p.add_run('关键词：')
            set_run_font(run1, cn_font='宋体', en_font='Times New Roman', size=12, bold=True)
            run2 = p.add_run(text)
            set_run_font(run2, cn_font='宋体', en_font='Times New Roman', size=12)
            i += 1
            continue

        # ===== 参考文献 =====
        if re.match(r'^\[\d+\]', line):
            add_ref_item(doc, line)
            i += 1
            continue

        # ===== 加粗行（**xxx**） =====
        text = line.strip()
        if text.startswith('**') and text.endswith('**') and text.count('**') == 2:
            inner = text[2:-2]
            p = doc.add_paragraph()
            set_paragraph_format(p, line_spacing=20, space_before=10, space_after=10)
            run = p.add_run(inner)
            set_run_font(run, cn_font='宋体', en_font='Times New Roman', size=12, bold=True)
            i += 1
            continue

        # ===== 普通正文 =====
        if text:
            add_body_text(doc, text)

        i += 1

    # 刷新剩余表格
    if table_buffer:
        rows = parse_table(table_buffer)
        add_table(doc, rows)

    doc.save(docx_path)
    print(f'Word文档已生成: {docx_path}')


if __name__ == '__main__':
    paper_dir = os.path.normpath(
        os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "论文")
    )
    md_file = os.path.join(paper_dir, '桂林气候舒适度与旅游影响分析论文.md')
    docx_file = os.path.join(paper_dir, '桂林气候舒适度与旅游影响分析论文.docx')
    md_to_docx(md_file, docx_file)
